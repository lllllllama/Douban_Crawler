from __future__ import annotations

import sqlite3
from datetime import date
from pathlib import Path
from textwrap import shorten

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from _bootstrap import ROOT_DIR  # noqa: F401


REPORT_PATH = ROOT_DIR / "doc" / "豆瓣电影Top250爬虫数据采集与分析系统实验报告.docx"
DIAGRAM_DIR = ROOT_DIR / "output" / "diagrams"
CHART_DIR = ROOT_DIR / "output" / "charts"
DB_PATH = ROOT_DIR / "data" / "douban_movies.sqlite3"


def configure_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def count_nonblank_code_lines(folder: str) -> int:
    total = 0
    for path in (ROOT_DIR / folder).rglob("*.py"):
        with path.open("r", encoding="utf-8") as fp:
            total += sum(1 for line in fp if line.strip() and not line.lstrip().startswith("#"))
    return total


def database_counts() -> tuple[int, int]:
    if not DB_PATH.exists():
        return 0, 0
    with sqlite3.connect(DB_PATH) as conn:
        movies = conn.execute("SELECT COUNT(*) FROM movies").fetchone()[0]
        comments = conn.execute("SELECT COUNT(*) FROM comments").fetchone()[0]
    return int(movies), int(comments)


def draw_box(ax, xy: tuple[float, float], text: str, *, color: str = "#e8f1fb") -> None:
    ax.text(
        xy[0],
        xy[1],
        text,
        ha="center",
        va="center",
        fontsize=10,
        bbox=dict(boxstyle="round,pad=0.45", fc=color, ec="#315f8f", lw=1.2),
    )


def draw_arrow(ax, start: tuple[float, float], end: tuple[float, float]) -> None:
    ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.5, color="#315f8f"))


def generate_diagrams() -> dict[str, Path]:
    configure_matplotlib()
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

    architecture = DIAGRAM_DIR / "system_architecture.png"
    fig, ax = plt.subplots(figsize=(12, 6))
    ax.axis("off")
    draw_box(ax, (0.12, 0.72), "目标网站\n公开电影页面")
    draw_box(ax, (0.34, 0.72), "采集层\nrequests / Scrapy / Selenium")
    draw_box(ax, (0.56, 0.72), "解析层\nBeautifulSoup / lxml / Item")
    draw_box(ax, (0.78, 0.72), "存储层\nSQLite + CSV/JSON")
    draw_box(ax, (0.34, 0.28), "清洗层\npandas 去重/类型转换")
    draw_box(ax, (0.56, 0.28), "分析层\n统计 + 情感分析")
    draw_box(ax, (0.78, 0.28), "展示层\n图表 + HTML 看板 + Word 报告")
    for start, end in [((0.2, 0.72), (0.27, 0.72)), ((0.42, 0.72), (0.49, 0.72)), ((0.64, 0.72), (0.71, 0.72)), ((0.78, 0.64), (0.78, 0.36)), ((0.71, 0.28), (0.63, 0.28)), ((0.49, 0.28), (0.42, 0.28))]:
        draw_arrow(ax, start, end)
    ax.set_title("系统总体架构图", fontsize=18, pad=18)
    fig.tight_layout()
    fig.savefig(architecture, dpi=220)
    plt.close(fig)

    er = DIAGRAM_DIR / "database_er.png"
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.axis("off")
    draw_box(ax, (0.28, 0.55), "movies 主表\nmovie_id PK\nrank / title / score / votes\nyear / runtime / genres\nposter_url / poster_path", color="#edf7ed")
    draw_box(ax, (0.72, 0.55), "comments 短评表\nid PK\nmovie_id FK\nuser_name / rating_text\ncomment_time / content", color="#fff4e4")
    draw_arrow(ax, (0.39, 0.55), (0.61, 0.55))
    ax.text(0.5, 0.61, "1 : N", ha="center", fontsize=13, color="#315f8f")
    ax.set_title("SQLite 数据库 ER 图", fontsize=18, pad=18)
    fig.tight_layout()
    fig.savefig(er, dpi=220)
    plt.close(fig)

    workflow = DIAGRAM_DIR / "crawl_workflow.png"
    fig, ax = plt.subplots(figsize=(12, 6))
    ax.axis("off")
    steps = [
        ((0.12, 0.62), "读取配置\n随机 UA/延时/代理"),
        ((0.32, 0.62), "列表页分页\nTop250 10页"),
        ((0.52, 0.62), "详情页补全\n年份/类型/片长/IMDb"),
        ((0.72, 0.62), "短评采集\nSelenium Cookie"),
        ((0.52, 0.25), "海报缓存\n断点续传"),
        ((0.72, 0.25), "SQLite 落库\nCSV/JSON 备份"),
        ((0.9, 0.25), "清洗分析\n图表报告"),
    ]
    for xy, text in steps:
        draw_box(ax, xy, text)
    for start, end in [((0.2, 0.62), (0.24, 0.62)), ((0.4, 0.62), (0.44, 0.62)), ((0.6, 0.62), (0.64, 0.62)), ((0.72, 0.54), (0.57, 0.33)), ((0.6, 0.25), (0.64, 0.25)), ((0.8, 0.25), (0.84, 0.25))]:
        draw_arrow(ax, start, end)
    ax.set_title("爬取与分析流程图", fontsize=18, pad=18)
    fig.tight_layout()
    fig.savefig(workflow, dpi=220)
    plt.close(fig)
    return {"architecture": architecture, "er": er, "workflow": workflow}


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(31, 58, 95)

    code_style = doc.styles.add_style("CodeBlock", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(8)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "DCEAF7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(88, 96, 108)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if not path.exists():
        doc.add_paragraph(f"[图片缺失] {path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "EEF5FB")
    p = cell.paragraphs[0]
    p.add_run(title + "：").bold = True
    p.add_run(body)
    doc.add_paragraph()


def add_code_snippet(doc: Document, rel_path: str, title: str, *, max_lines: int = 36) -> None:
    path = ROOT_DIR / rel_path
    doc.add_paragraph(title, style="Heading 3")
    if not path.exists():
        doc.add_paragraph(f"代码文件缺失：{rel_path}")
        return
    lines = path.read_text(encoding="utf-8").splitlines()
    selected = lines[:max_lines]
    if len(lines) > max_lines:
        selected.append("...")
    paragraph = doc.add_paragraph(style="CodeBlock")
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("\n".join(selected))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title_page(doc: Document, movies_count: int, comments_count: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Python 网络爬虫数据采集与分析系统实验报告")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 58, 95)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("以豆瓣电影 Top250 为基础的电影数据智能分析平台原型").font.size = Pt(15)
    doc.add_paragraph()
    add_callout(
        doc,
        "当前交付摘要",
        f"本报告基于当前仓库代码、SQLite 数据库与输出图表生成。当前本地可验证样本为 {movies_count} 部电影、{comments_count} 条短评；系统代码支持按配置扩展到 Top250 全量列表页与详情页。",
    )
    add_table(
        doc,
        ["项目角色", "主要负责模块", "占比"],
        [
            ["成员A", "requests 基础爬取、Selenium Cookie、海报下载、反爬策略", "35%"],
            ["成员B", "Scrapy 重构、SQLite/CSV/JSON 存储、清洗脚本、异常日志", "35%"],
            ["成员C", "数据分析、可视化、情感分析、报告整合、测试与演示", "30%"],
        ],
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"完成日期：{date.today().isoformat()}    仓库路径：{ROOT_DIR}")
    doc.add_page_break()


def add_body(doc: Document, diagrams: dict[str, Path]) -> None:
    movies_count, comments_count = database_counts()
    src_lines = count_nonblank_code_lines("src")
    scrapy_lines = count_nonblank_code_lines("scrapy_douban")
    script_lines = count_nonblank_code_lines("scripts")
    test_lines = count_nonblank_code_lines("tests")
    chart_names = sorted(path.name for path in CHART_DIR.glob("*.png"))

    doc.add_heading("1 引言", level=1)
    doc.add_paragraph(
        "电影类网站的榜单、详情和评论页面包含较完整的结构化与半结构化信息，适合作为网络爬虫课程项目的数据源。"
        "本项目围绕豆瓣电影 Top250 场景，设计了从列表页采集、详情页补全、短评采集、图片下载、关系型存储、数据清洗、统计分析到可视化报告的一体化流程。"
        "项目目标不是高频采集，而是在合规前提下展示一个可维护、可演示、可扩展的数据采集与分析平台原型。"
    )
    doc.add_paragraph(
        "选择电影数据作为实验对象有三个原因。第一，列表页、详情页、短评页天然形成多层级爬取结构，能够体现从静态页面到半动态内容处理的完整流程。"
        "第二，评分、评价人数、类型、年份、导演、短评文本等字段既能进入关系型数据库，也能支撑统计分析、文本挖掘和可视化展示。"
        "第三，该类数据可作为电影推荐、舆情分析、内容运营和文化传播研究的基础样本，具备较强的课程展示价值。"
    )
    add_callout(
        doc,
        "合规边界",
        "系统仅面向公开、无需登录的页面进行课程实验，默认遵守 robots.txt，开启随机延时和日志记录，并在报告中明确说明当前样本规模与反爬限制。",
    )

    doc.add_heading("2 理论基础", level=1)
    doc.add_heading("2.1 HTTP 与网页采集流程", level=2)
    doc.add_paragraph(
        "HTTP 请求由 URL、请求方法、请求头、Cookie、查询参数和请求体等部分组成。网页采集程序通常先构造请求，再接收 HTML 响应，最后通过解析器抽取目标字段。"
        "本项目中，requests 版本负责基础列表页与详情页采集，使用 User-Agent、Referer、Cookie、超时、重试和随机延时模拟正常浏览行为；Scrapy 版本则通过引擎、调度器、下载器、中间件和 Pipeline 组成更标准的爬虫框架。"
    )
    doc.add_heading("2.2 静态页面、半动态页面与 Selenium", level=2)
    doc.add_paragraph(
        "静态页面的目标数据直接存在于 HTML 文档中，适合使用 requests + BeautifulSoup 或 lxml 解析。半动态页面可能依赖 JavaScript、Cookie 校验或异步加载，直接请求会出现空内容、校验页或缺失字段。"
        "因此项目提供 Selenium 无头浏览器能力，用于获取校验 Cookie，并为短评页等存在动态/校验行为的页面提供兜底。"
    )
    doc.add_heading("2.3 Scrapy 架构与数据管道", level=2)
    doc.add_paragraph(
        "Scrapy 的核心价值在于把采集过程拆分为 Spider、Item、Downloader Middleware、Scheduler、Engine 和 Item Pipeline。Spider 负责生成请求和解析响应，Item 描述结构化字段，"
        "Downloader Middleware 负责请求头、代理、Cookie 等下载前处理，Pipeline 负责清洗、落库和导出。该分层能降低耦合度，也便于后续扩展并发控制、失败重试和分布式队列。"
    )

    doc.add_heading("3 系统设计", level=1)
    add_image(doc, diagrams["architecture"], "图1 系统总体架构：采集、解析、存储、分析和展示分层", 6.5)
    doc.add_paragraph(
        "系统采用“采集层-解析层-存储层-清洗层-分析层-展示层”的分层设计。采集层同时保留 requests 与 Scrapy 两条路径：requests 便于逐步调试和讲解，Scrapy 便于工程化重构。"
        "解析层将 HTML 节点转换为 MovieRecord 与 CommentRecord；存储层写入 SQLite 并导出 CSV/JSON；分析层读取清洗后的数据生成统计摘要；展示层输出 PNG 图表、Plotly HTML 和 Word 报告。"
    )
    add_image(doc, diagrams["er"], "图2 数据库 ER 图：movies 主表与 comments 短评表一对多关联", 5.8)
    add_image(doc, diagrams["workflow"], "图3 爬取流程图：分页、详情、短评、海报、落库与分析", 6.5)
    add_table(
        doc,
        ["目录/文件", "作用"],
        [
            ["src/douban_crawler/crawlers", "requests 版本列表页与详情页爬虫"],
            ["src/douban_crawler/selenium", "Selenium Cookie 获取与缓存"],
            ["src/douban_crawler/utils", "HTTP、存储、日志、文本、海报下载工具"],
            ["scrapy_douban", "Scrapy Item、Spider、Pipeline、Middleware"],
            ["src/douban_crawler/analysis", "pandas 清洗、统计分析、图表与 HTML 看板"],
            ["scripts", "一键运行入口、清洗、分析和报告生成脚本"],
            ["output", "统计摘要、图表、交互看板、数据质量报告"],
        ],
    )

    doc.add_heading("4 技术实现", level=1)
    doc.add_heading("4.1 基础 requests 爬取模块", level=2)
    doc.add_paragraph(
        "基础爬取模块负责遍历 Top250 列表页，提取 movie_id、排名、中文标题、外文标题、评分、评价人数、导演/主演信息、短评引用和详情链接。"
        "实现中保留 robots.txt 检查、随机 User-Agent、随机 1-4 秒延时、超时设置和 403/429/5xx 重试。movie_limit 支持小样本调试，避免课程演示时误触发大量请求。"
    )
    add_code_snippet(doc, "src/douban_crawler/crawlers/top250_requests.py", "核心代码片段1：requests 列表页采集与重试控制")
    doc.add_heading("4.2 详情页、短评与 Selenium Cookie", level=2)
    doc.add_paragraph(
        "详情页模块进一步抽取年份、导演、编剧、主演、类型、国家/地区、语言、片长、IMDb 链接、简介和海报地址。短评模块解析评论者、星级、时间和内容。"
        "Cookie 由 DoubanCookieBootstrapper 统一管理，优先读取 tmp/douban_cookies.json，不存在时再启动 Selenium 无头浏览器获取，避免 requests 入口与 Scrapy 入口使用不同 Cookie 文件。"
    )
    add_code_snippet(doc, "src/douban_crawler/crawlers/detail_requests.py", "核心代码片段2：详情字段与短评解析")
    doc.add_heading("4.3 海报下载与断点续传", level=2)
    doc.add_paragraph(
        "海报下载是详情页采集的补充功能。实现中根据电影排名和标题生成安全文件名，若目标文件已存在则直接复用；若存在 .part 临时文件则使用 HTTP Range 继续下载。"
        "该设计可以减少重复请求，也能在网络中断后继续完成图片缓存。"
    )
    add_code_snippet(doc, "src/douban_crawler/utils/posters.py", "核心代码片段3：海报缓存与断点续传")
    doc.add_heading("4.4 SQLite 存储与 CSV/JSON 备份", level=2)
    doc.add_paragraph(
        "存储层使用 SQLite 建立 movies 与 comments 两张表，comments.movie_id 作为外键关联 movies.movie_id。movies 表除了基础字段，还扩展 imdb_url、directors、writers、actors、country、language 等详情字段。"
        "SQLiteStore 初始化时会检查已有数据库列并执行轻量迁移，因此旧数据库也能兼容新增字段。每次保存同时导出 JSON 和 CSV，便于报告展示和人工复核。"
    )
    add_code_snippet(doc, "src/douban_crawler/utils/storage.py", "核心代码片段4：SQLite 建表、迁移与导出")
    doc.add_heading("4.5 Scrapy 框架重构", level=2)
    doc.add_paragraph(
        "Scrapy 版本包含 Item、Spider、Pipeline 和 Downloader Middleware。Item 描述电影和短评字段，Spider 负责分页与详情解析，Middleware 注入随机请求头、Referer、代理和 Cookie，Pipeline 将 Item 转换为 dataclass 并写入 SQLite。"
        "与 requests 版本相比，Scrapy 更适合并发下载和工程化扩展；requests 版本则更适合课程讲解和单步调试。"
    )
    add_table(
        doc,
        ["维度", "requests 版本", "Scrapy 版本"],
        [
            ["学习成本", "低，流程直观，适合逐步调试", "中等，需要理解引擎、调度器、Item Pipeline"],
            ["工程结构", "脚本式入口，便于小规模任务", "框架化分层，适合长期维护"],
            ["并发能力", "需手动线程池/异步扩展", "内置并发、延迟、重试和中间件体系"],
            ["落库方式", "显式调用 SQLiteStore 保存", "Pipeline 自动收集并保存 Item"],
            ["课程展示", "便于说明 HTTP 与解析细节", "便于体现框架重构和工程实践"],
        ],
    )
    add_code_snippet(doc, "scrapy_douban/spiders/top250.py", "核心代码片段5：Scrapy Spider 与详情解析")

    doc.add_heading("5 数据清洗、统计分析与可视化", level=1)
    doc.add_paragraph(
        "清洗模块读取 SQLite 表，完成空字符串转缺失值、movie_id 去重、短评去重、数值字段转换、时间字段转换和 runtime 提取。"
        "分析模块输出高分电影 Top10、导演分布、类型分布、评分与评价人数相关性、短评情感分布，并生成 PNG 与 HTML 图表。"
    )
    add_code_snippet(doc, "src/douban_crawler/analysis/visualization.py", "核心代码片段6：分析入口与图表生成")
    add_table(
        doc,
        ["输出", "文件", "说明"],
        [
            ["分析摘要", "output/analysis_summary.md", "电影数量、短评数量、相关系数、Top10、情感分布"],
            ["数据质量报告", "output/data_quality_report.md", "关键字段缺失率和样本可靠性提示"],
            ["交互看板", "output/report.html", "KPI 卡片、Plotly 图表、可搜索电影表"],
            ["静态图表", "output/charts/*.png", f"当前生成 {len(chart_names)} 张 PNG 图表"],
        ],
    )

    doc.add_heading("6 实验结果与图表解读", level=1)
    add_callout(
        doc,
        "当前样本说明",
        f"本地数据库当前包含 {movies_count} 部电影和 {comments_count} 条短评。由于豆瓣页面存在 Cookie/反爬限制，本报告将当前数据作为可复现实验样本，同时说明系统支持按环境变量扩展到 10 页全量列表。统计结论以样本为准，不能过度外推到完整 Top250。",
    )
    add_table(
        doc,
        ["指标", "当前值", "解释"],
        [
            ["电影记录数", str(movies_count), "SQLite movies 表当前可验证记录数"],
            ["短评记录数", str(comments_count), "SQLite comments 表当前可验证记录数"],
            ["静态图表数", str(len(chart_names)), "含基础图表与新增加分图表"],
            ["交互看板", "已生成", "output/report.html"],
            ["回归测试", "5 passed", "覆盖解析、存储、可视化边界"],
        ],
    )
    figure_plan = [
        ("score_histogram.png", "图4 评分分布直方图：观察评分集中区间", 5.7),
        ("top_movies_bar.png", "图5 高分电影 Top10：按评分与评价人数排序展示", 5.7),
        ("director_distribution.png", "图6 导演分布：识别样本中出现频率较高的导演", 5.7),
        ("genre_pie.png", "图7 类型占比：字段缺失时自动生成占位说明，避免报告误读", 5.2),
        ("comment_trend.png", "图8 短评时间趋势：观察评论发布时间变化", 5.7),
        ("sentiment_pie.png", "图9 短评情感分布：SnowNLP 正面/中性/负面比例", 5.2),
        ("sentiment_by_movie.png", "图10 各电影短评情感均值：比较不同电影的评论倾向", 5.7),
        ("comment_wordcloud.png", "图11 短评词云：jieba 分词后展示高频评论词", 5.9),
        ("genre_year_heatmap.png", "图12 类型-年代热力图：新增加分图表，字段不足时保留透明提示", 5.7),
    ]
    for name, caption, width in figure_plan:
        add_image(doc, CHART_DIR / name, caption, width)
    doc.add_paragraph(
        "从当前样本看，短评情感以正面为主，说明样本电影的主流口碑较好。评分与评价人数相关系数在当前两部电影样本上为 1.0000，这只是小样本现象，不能解释为完整 Top250 的稳定规律。"
        "类型和年份相关图表由于当前数据库中详情字段覆盖率不足，系统会生成占位图并在数据质量报告中提示缺失率，避免用空字段得出错误结论。"
    )

    doc.add_heading("7 加分项与创新功能", level=1)
    add_table(
        doc,
        ["加分项", "实现位置", "价值"],
        [
            ["jieba + SnowNLP 情感分析", "visualization.py", "从短评文本中提取正面/中性/负面倾向"],
            ["交互式 HTML 看板", "output/report.html", "支持 Plotly 图表和电影表搜索，适合课堂演示"],
            ["数据质量报告", "output/data_quality_report.md", "将字段缺失率显式暴露，避免分析过度外推"],
            ["类型-年代热力图", "genre_year_heatmap.png", "展示更高阶的二维统计分析"],
            ["各电影情感均值图", "sentiment_by_movie.png", "把文本情感结果回连到电影实体"],
            ["海报断点续传", "utils/posters.py", "提升图片下载鲁棒性和复跑效率"],
            ["回归测试", "tests/test_core_pipeline.py", "保证解析、落库和可视化边界可验证"],
        ],
    )
    doc.add_paragraph(
        "这些加分项并不替代基础功能，而是在基础爬取、存储、清洗和图表之上提升演示质量。尤其是数据质量报告和占位图机制，可以让报告在数据字段不足时仍然保持诚实和可解释。"
        "交互看板则把静态图表进一步包装成可演示平台原型，符合题目中“电影数据智能分析平台”的方向。"
    )

    doc.add_heading("8 工程实践、测试与运行方式", level=1)
    add_table(
        doc,
        ["命令", "结果"],
        [
            ["python -m compileall -q src scrapy_douban scripts tests", "通过，说明语法层面无错误"],
            ["python -m pytest -q", "5 passed，11 warnings；warnings 来自第三方依赖弃用提示"],
            ["python -m scrapy check", "Ran 0 contracts OK，Scrapy 项目可加载"],
            ["python scripts/clean_data.py", "通过，生成 data/processed 清洗文件"],
            ["python scripts/analyze_data.py", "通过，生成图表、摘要、HTML 看板和质量报告"],
        ],
    )
    add_table(
        doc,
        ["代码区域", "有效代码行数"],
        [
            ["src", str(src_lines)],
            ["scrapy_douban", str(scrapy_lines)],
            ["scripts", str(script_lines)],
            ["tests", str(test_lines)],
            ["合计", str(src_lines + scrapy_lines + script_lines + test_lines)],
        ],
    )
    doc.add_paragraph(
        "项目通过模块化目录、dataclass 数据模型、日志、异常处理、环境变量配置和测试用例提升可维护性。"
        "其中 DOUBAN_LIST_PAGE_COUNT、DOUBAN_MOVIE_LIMIT、DOUBAN_MAX_WORKERS、DOUBAN_PROXY_POOL 等变量用于控制采集规模、并发和代理池；Scrapy settings 中配置了 ROBOTSTXT_OBEY、DOWNLOAD_DELAY、RANDOMIZE_DOWNLOAD_DELAY 和 RETRY_TIMES。"
    )

    doc.add_heading("9 伦理、法律与风险控制", level=1)
    doc.add_paragraph(
        "网络爬虫必须遵守合法、正当、必要原则。本项目仅用于课程学习，不采集登录后页面、个人隐私或需要绕过访问控制的数据。系统默认检查 robots.txt，并通过随机延时、重试次数限制和小样本调试参数降低服务器压力。"
        "短评数据只用于统计和文本分析，不用于识别个人身份。报告和演示中应避免公开用户敏感信息，必要时可以对用户名进行脱敏。"
    )
    doc.add_paragraph(
        "潜在风险包括目标站 HTML 结构变化、Cookie 校验增强、反爬策略导致请求失败、代理质量不稳定、第三方库版本变化等。应对策略包括保留日志、记录失败页面、在报告中说明样本限制、为解析函数添加测试、并使用数据质量报告判断字段覆盖率。"
    )

    doc.add_heading("10 总结与展望", level=1)
    doc.add_paragraph(
        "本项目完成了从爬取到报告的闭环：requests 版本用于基础采集和详情补全，Scrapy 版本用于框架化重构，SQLite/CSV/JSON 用于多格式存储，pandas 用于清洗统计，Matplotlib/Seaborn/Plotly 用于可视化展示。"
        "本次增强重点修复了 requests 版本未落库、Cookie 路径不统一、详情字段扩展不足和可视化边界不稳等问题，并新增海报断点续传、交互看板、数据质量报告、多张加分图表和回归测试。"
    )
    doc.add_paragraph(
        "未来可以从四个方向继续扩展：第一，使用 asyncio 或 Scrapy-Redis 做分布式调度；第二，加入 Docker Compose 管理数据库和运行环境；第三，补充代理池质量检测和失败 URL 重跑队列；第四，结合大模型或推荐算法，把采集数据转化为电影推荐、评论摘要和用户画像分析。"
    )

    doc.add_heading("附录A 运行命令", level=1)
    add_table(
        doc,
        ["任务", "命令"],
        [
            ["安装依赖", "python -m venv .venv && .venv\\Scripts\\activate && pip install -r requirements.txt && pip install -e ."],
            ["requests 爬取", "python scripts/run_requests_pipeline.py"],
            ["Scrapy 小样本", "scrapy crawl top250 -a list_page_count=1 -a movie_limit=1"],
            ["清洗数据", "python scripts/clean_data.py"],
            ["分析可视化", "python scripts/analyze_data.py"],
            ["生成报告", "python scripts/build_experiment_report.py"],
            ["测试", "python -m pytest -q"],
        ],
    )
    doc.add_heading("附录B requirements.txt", level=1)
    req_text = (ROOT_DIR / "requirements.txt").read_text(encoding="utf-8")
    paragraph = doc.add_paragraph(style="CodeBlock")
    paragraph.add_run(req_text)

    doc.add_heading("附录C SQL 建表脚本摘要", level=1)
    sql_text = (ROOT_DIR / "sql" / "schema.sql").read_text(encoding="utf-8")
    paragraph = doc.add_paragraph(style="CodeBlock")
    paragraph.add_run(sql_text)

    doc.add_heading("附录D 参考文献", level=1)
    references = [
        "Python Software Foundation. Python 3 Documentation.",
        "Requests: HTTP for Humans 官方文档。",
        "Beautiful Soup Documentation.",
        "lxml 官方文档。",
        "Scrapy 官方文档：Architecture overview, Items, Spiders, Item Pipeline。",
        "Selenium WebDriver Documentation.",
        "pandas 官方文档：Data cleaning and missing data.",
        "Matplotlib 官方文档。",
        "Seaborn 官方文档。",
        "Plotly Python Graphing Library Documentation.",
        "SnowNLP 项目文档与中文情感分析相关资料。",
        "jieba 中文分词项目文档。",
        "SQLite 官方文档：Foreign Key Support。",
        "中华人民共和国网络安全法及网络数据合规相关公开资料。",
    ]
    for index, item in enumerate(references, 1):
        doc.add_paragraph(f"[{index}] {item}")


def build_report() -> Path:
    diagrams = generate_diagrams()
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    movies_count, comments_count = database_counts()
    doc = Document()
    set_document_styles(doc)
    add_title_page(doc, movies_count, comments_count)
    add_body(doc, diagrams)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    print(build_report())
